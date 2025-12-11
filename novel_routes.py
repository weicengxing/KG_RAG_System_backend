"""
小说阅读系统路由
提供小说列表、内容获取、VIP状态管理等API接口
"""

import os
import logging
import time
import hashlib
import uuid
from pathlib import Path
from typing import Optional
from fastapi import APIRouter, HTTPException, Query, Depends, UploadFile, File, Form, Request
from fastapi.responses import JSONResponse
from database import driver
from database_async import AsyncDatabaseManager
from redis_utils import RedisDistributedLock, acquire_distributed_lock
from config import NEO4J_URI, NEO4J_USERNAME, NEO4J_PASSWORD

# 配置日志
logger = logging.getLogger(__name__)

# 创建路由
router = APIRouter(prefix="/deal", tags=["小说"])

# 小说目录
NOVELS_DIR = Path("D:/创作")

# 异步数据库管理器
db_async_manager = AsyncDatabaseManager(max_workers=5)

# 导入docx处理库（如果可用）
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 库未安装，将无法处理 .doc/.docx 文件")


def get_novels_from_filesystem():
    """从文件系统获取小说列表"""
    novels = []
    if not NOVELS_DIR.exists():
        logger.error(f"小说目录不存在: {NOVELS_DIR}")
        return novels

    txt_files = sorted(NOVELS_DIR.glob("*.txt"))
    for idx, file_path in enumerate(txt_files, 1):
        title = file_path.stem  # 文件名（不含扩展名）
        # 简单规则：偶数序号的小说为VIP专属
        quan = 1 if idx % 2 == 1 else 0
        novels.append({
            "id": idx,
            "title": title,
            "path": str(file_path),
            "quan": quan
        })
    return novels


def get_novels_from_neo4j():
    """从Neo4j获取小说列表"""
    try:
        with driver.session() as session:
            result = session.run("""
                MATCH (n:Novel)
                RETURN n.id as id, n.title as title, n.path as path, n.quan as quan
                ORDER BY n.id
            """)
            novels = []
            for record in result:
                novels.append({
                    "id": record["id"],
                    "title": record["title"],
                    "path": record["path"],
                    "quan": record["quan"]
                })
            if novels:
                return novels
    except Exception as e:
        logger.warning(f"从Neo4j获取小说列表失败: {e}")

    return get_novels_from_filesystem()


@router.post("/getnovelzong")
async def get_novel_list():
    """获取所有小说列表"""
    novels = get_novels_from_neo4j()
    return [{"title": n["title"], "path": n["path"], "quan": n["quan"]} for n in novels]


@router.post("/getnovel")
async def get_novel_content(
    a: int = Query(..., description="章节号/小说序号"),
    b: str = Query("", description="VIP状态标识，'1'表示VIP")
):
    """获取小说内容"""
    novels = get_novels_from_neo4j()

    if a < 1 or a > len(novels):
        raise HTTPException(status_code=404, detail="小说不存在")

    novel = novels[a - 1]
    is_vip = b == "1"

    try:
        file_path = Path(novel["path"])
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="小说文件不存在")

        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        # 如果是VIP专属小说且用户不是VIP，只返回40%内容
        if novel["quan"] == 0 and not is_vip:
            content_length = len(content)
            visible_length = int(content_length * 0.4)
            content = content[:visible_length]

        return content

    except Exception as e:
        logger.error(f"读取小说内容失败: {e}")
        raise HTTPException(status_code=500, detail="读取小说内容失败")


@router.post("/gethui")
async def get_vip_status(b: str = Query("", description="用户名")):
    """获取用户VIP状态"""
    if not b:
        return ""

    try:
        with driver.session() as session:
            result = session.run("""
                MATCH (u:User {username: $username})
                RETURN u.is_vip as is_vip
            """, username=b)
            record = result.single()
            if record and record["is_vip"]:
                return "1"
    except Exception as e:
        logger.warning(f"从Neo4j获取VIP状态失败: {e}")

    return ""


@router.post("/setvip")
async def set_vip_status(
    username: str = Query(..., description="用户名"),
    is_vip: bool = Query(..., description="是否VIP")
):
    """设置用户VIP状态"""
    try:
        with driver.session() as session:
            session.run("""
                MATCH (u:User {username: $username})
                SET u.is_vip = $is_vip
            """, username=username, is_vip=is_vip)
        return {"success": True, "username": username, "is_vip": is_vip}
    except Exception as e:
        logger.error(f"保存VIP状态到Neo4j失败: {e}")
        raise HTTPException(status_code=500, detail="VIP状态更新失败")


@router.get("/getvip")
async def get_user_vip(username: str = Query(..., description="用户名")):
    """获取用户VIP状态（GET方式）"""
    is_vip = False

    try:
        with driver.session() as session:
            result = session.run("""
                MATCH (u:User {username: $username})
                RETURN u.is_vip as is_vip
            """, username=username)
            record = result.single()
            if record and record["is_vip"] is not None:
                is_vip = record["is_vip"]
    except Exception as e:
        logger.warning(f"从Neo4j获取VIP状态失败: {e}")

    return {"is_vip": is_vip, "username": username}


# ==================== VIP用户上传小说相关功能 ====================

def get_current_user(request: Request) -> Optional[str]:
    """从请求中获取当前登录用户
    
    Args:
        request: FastAPI请求对象
        
    Returns:
        Optional[str]: 用户名，如果未登录返回None
    """
    username = getattr(request.state, "current_user", None)
    return username


def check_user_vip_status(username: str) -> bool:
    """检查用户是否为VIP
    
    Args:
        username: 用户名
        
    Returns:
        bool: 是否为VIP
    """
    try:
        with driver.session() as session:
            result = session.run("""
                MATCH (u:User {username: $username})
                RETURN u.is_vip as is_vip
            """, username=username)
            record = result.single()
            if record and record["is_vip"]:
                return True
    except Exception as e:
        logger.error(f"检查用户VIP状态失败: {username}, 错误: {e}")
    
    return False


def extract_text_from_docx(file_content: bytes) -> str:
    """从.docx文件中提取文本内容
    
    Args:
        file_content: 文件内容字节
        
    Returns:
        str: 提取的文本内容
    """
    if not DOCX_AVAILABLE:
        raise ValueError("python-docx 库未安装，无法处理 .docx 文件")
    
    import io
    from docx import Document
    
    try:
        # 创建内存文件对象
        file_stream = io.BytesIO(file_content)
        doc = Document(file_stream)
        
        # 提取所有段落文本，并用换行符连接
        # 优化：同时提取表格内容，防止内容丢失
        text_content = []
        
        # 提取段落
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 提取表格（可选，视需求而定，这里简单提取）
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content)
    except Exception as e:
        logger.error(f"解析docx文件失败: {e}")
        return "" # 解析失败返回空字符串，由上层处理


def get_max_novel_id() -> int:
    """获取当前最大小说ID
    
    Returns:
        int: 最大ID，如果没有小说返回0
    """
    try:
        with driver.session() as session:
            result = session.run("""
                MATCH (n:Novel)
                RETURN max(n.id) as max_id
            """)
            record = result.single()
            if record and record["max_id"] is not None:
                return int(record["max_id"])
    except Exception as e:
        logger.error(f"获取最大小说ID失败: {e}")
    
    # 如果数据库中没有小说，检查文件系统
    txt_files = list(NOVELS_DIR.glob("*.txt"))
    return len(txt_files)


def save_novel_file_sync(title: str, content: str, file_extension: str = ".txt") -> str:
    """同步保存小说文件到文件系统
    
    Args:
        title: 小说标题
        content: 小说内容
        file_extension: 文件扩展名，默认.txt
        
    Returns:
        str: 保存的文件路径
    """
    # 确保目录存在
    NOVELS_DIR.mkdir(parents=True, exist_ok=True)
    
    # 清理文件名（移除非法字符）
    clean_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).strip()
    if not clean_title:
        clean_title = "未命名小说"
    
    # 生成文件名（防止重复）
    timestamp = int(time.time() * 1000)
    unique_id = uuid.uuid4().hex[:8]
    filename = f"{clean_title}_{timestamp}_{unique_id}{file_extension}"
    file_path = NOVELS_DIR / filename
    
    # 保存文件
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(content)
    
    logger.info(f"小说文件已保存: {file_path}")
    return str(file_path)


def save_novel_to_db_sync(novel_data: dict) -> bool:
    """同步保存小说信息到数据库
    
    Args:
        novel_data: 小说数据字典，包含id, title, path, quan
        
    Returns:
        bool: 是否保存成功
    """
    try:
        with driver.session() as session:
            session.run("""
                CREATE (n:Novel {
                    id: $id,
                    title: $title,
                    path: $path,
                    quan: $quan,
                    created_at: timestamp(),
                    uploader: $uploader
                })
            """, 
            id=novel_data["id"],
            title=novel_data["title"],
            path=novel_data["path"],
            quan=novel_data["quan"],
            uploader=novel_data.get("uploader", "")
            )
        logger.info(f"小说信息已保存到数据库: ID={novel_data['id']}, 标题={novel_data['title']}")
        return True
    except Exception as e:
        logger.error(f"保存小说信息到数据库失败: {e}")
        return False


async def async_save_novel(title: str, content: str, quan: int, uploader: str) -> dict:
    """异步保存小说（文件系统 + 数据库）
    
    使用分布式锁确保ID生成的原子性
    
    Args:
        title: 小说标题
        content: 小说内容
        quan: VIP状态（0=VIP, 1=非VIP）
        uploader: 上传者用户名
        
    Returns:
        dict: 保存结果，包含成功状态和小说信息
    """
    # 获取分布式锁
    lock = acquire_distributed_lock("novel_upload_id_generation", expire_time=30, timeout=10)
    if not lock:
        return {"success": False, "error": "系统繁忙，请稍后重试"}
    
    try:
        # 在锁的保护下获取最大ID并生成新ID
        max_id = get_max_novel_id()
        new_id = max_id + 1
        
        # 保存文件到文件系统（同步，因为需要在获取锁期间完成）
        file_path = save_novel_file_sync(title, content)
        
        # 准备小说数据
        novel_data = {
            "id": new_id,
            "title": title,
            "path": file_path,
            "quan": quan,
            "uploader": uploader
        }
        
        # 异步保存到数据库（不阻塞）
        success = await db_async_manager.submit_novel_save(novel_data)
        
        if success:
            return {
                "success": True,
                "novel_id": new_id,
                "title": title,
                "file_path": file_path,
                "quan": quan
            }
        else:
            # 如果数据库保存失败，删除文件
            try:
                Path(file_path).unlink()
            except Exception as e:
                logger.error(f"删除小说文件失败: {file_path}, 错误: {e}")
            
            return {"success": False, "error": "保存小说信息失败"}
            
    except Exception as e:
        logger.error(f"保存小说过程中发生错误: {e}")
        return {"success": False, "error": str(e)}
    finally:
        # 释放锁
        lock.release()


@router.post("/upload-novel")
async def upload_novel(
    request: Request,
    title: str = Form(..., description="小说标题"),
    quan: int = Form(1, description="VIP状态（0=VIP章节, 1=非VIP章节）"),
    content: Optional[str] = Form(None, description="文本内容（如果使用文本区域上传）"),
    file: Optional[UploadFile] = File(None, description="小说文件（支持.txt/.doc/.docx）")
):
    """VIP用户上传小说
    
    支持两种上传方式：
    1. 文本区域上传：直接提供文本内容
    2. 文件上传：上传.txt/.doc/.docx文件
    
    只有VIP用户可以使用此功能
    """
    # 获取当前用户
    username = get_current_user(request)
    if not username:
        raise HTTPException(status_code=401, detail="请先登录")
    
    # 检查用户是否为VIP
    if not check_user_vip_status(username):
        raise HTTPException(status_code=403, detail="此功能仅限VIP用户使用")
    
    # 验证输入
    if not title.strip():
        raise HTTPException(status_code=400, detail="小说标题不能为空")
    
    if quan not in (0, 1):
        raise HTTPException(status_code=400, detail="VIP状态参数无效（必须为0或1）")
    
    # 获取小说内容
    novel_content = ""
    file_ext = ".txt" # 默认为txt，如果是文件上传则覆盖
    
    if content and content.strip():
        # 使用文本区域内容
        novel_content = content.strip()
    elif file and file.filename:
        # 使用上传的文件
        # 检查文件类型
        filename_lower = file.filename.lower()
        file_ext = os.path.splitext(filename_lower)[1]
        
        if file_ext not in ('.txt', '.doc', '.docx'):
            raise HTTPException(status_code=400, detail="只支持.txt/.doc/.docx文件格式")
        
        # 读取文件内容
        file_content = await file.read()
        if len(file_content) > 10 * 1024 * 1024:  # 10MB限制
            raise HTTPException(status_code=400, detail="文件大小不能超过10MB")
        
        # 根据文件类型提取文本
        if file_ext == '.txt':
            # 尝试不同的编码格式
            try:
                novel_content = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    novel_content = file_content.decode('gbk')
                except UnicodeDecodeError:
                    # 如果都失败，使用 errors='ignore'
                    novel_content = file_content.decode('utf-8', errors='ignore')
                    
        elif file_ext == '.docx':
            if not DOCX_AVAILABLE:
                raise HTTPException(status_code=400, detail="系统暂时不支持.docx文件解析")
            try:
                extracted_text = extract_text_from_docx(file_content)
                if extracted_text:
                    novel_content = extracted_text
                else:
                    novel_content = "（该.docx文件内容解析为空或格式不支持在线预览，请下载阅读）"
            except Exception as e:
                logger.error(f"解析docx失败: {e}")
                novel_content = "（文件解析失败，请下载阅读）"
                
        elif file_ext == '.doc':
            # .doc 文件无法通过 python-docx 解析
            # 这里我们只保存文件，内容字段存入提示信息
            novel_content = "（.doc格式不支持在线预览，请下载文件阅读）"
            
    else:
        raise HTTPException(status_code=400, detail="请提供小说内容或上传文件")
    
    # 即使内容为空（例如解析失败），只要是文件上传，也允许保存文件
    # 但如果是纯文本上传且为空，则报错
    if not novel_content.strip() and not file:
        raise HTTPException(status_code=400, detail="小说内容不能为空")
    
    # 如果内容为空但有文件，设置默认提示
    if not novel_content.strip():
        novel_content = "（内容暂不可见）"

    # 异步保存小说
    # 注意：async_save_novel 会将 novel_content 保存到 .txt 文件中
    # 上传的是 .doc/.docx，实际上是将解析后的文本（或者提示语）保存为了 .txt
    # 原文件并没有被保存
    #
  
    
    result = await async_save_novel(title, novel_content, quan, username)
    
    if result["success"]:
        return JSONResponse(
            status_code=200,
            content={
                "success": True,
                "message": "小说上传成功",
                "novel_id": result["novel_id"],
                "title": result["title"],
                "file_path": result["file_path"],
                "quan": result["quan"]
            }
        )
    else:
        raise HTTPException(status_code=500, detail=result.get("error", "小说上传失败"))
    

@router.get("/check-upload-permission")
async def check_upload_permission(request: Request):
    """检查用户是否有上传小说的权限（用于前端按钮显示）
    
    Returns:
        dict: 包含权限信息和VIP状态
    """
    username = get_current_user(request)
    if not username:
        return {"has_permission": False, "is_vip": False, "message": "请先登录"}
    
    is_vip = check_user_vip_status(username)
    return {
        "has_permission": is_vip,
        "is_vip": is_vip,
        "username": username,
        "message": "VIP用户" if is_vip else "普通用户"
    }
